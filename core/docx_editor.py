import datetime
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class DocxEditor:
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = Document(file_path)
        self._enable_tracking()

    def _enable_tracking(self):
        """ドキュメント全体の変更履歴記録を有効にします。"""
        settings = self.doc.settings.element
        track_revisions = settings.find(qn('w:trackRevisions'))
        if track_revisions is None:
            track_revisions = OxmlElement('w:trackRevisions')
            settings.append(track_revisions)

    def get_text(self) -> str:
        """ドキュメント全体のテキストを取得します。"""
        return "\n".join([para.text for para in self.doc.paragraphs])

    def apply_corrections(self, corrections: list):
        """
        添削内容を適用し、インラインで変更履歴とコメントを挿入します。
        """
        for correction in corrections:
            original = correction.get("original")
            corrected = correction.get("corrected")
            reason = correction.get("reason")
            
            if not original:
                continue
                
            for paragraph in self.doc.paragraphs:
                # 段落内で複数回出現する場合に対応
                while original in paragraph.text:
                    parts = paragraph.text.split(original, 1)
                    if len(parts) == 2:
                        before, after = parts
                        paragraph.clear()
                        
                        # 置換前のテキスト
                        if before:
                            paragraph.add_run(before)
                            
                        # 変更履歴（削除と挿入）の追加
                        self._insert_revision(paragraph, original, corrected)
                        
                        # 置換後のテキスト
                        if after:
                            paragraph.add_run(after)
                        
                        # コメントの付与
                        if paragraph.runs and reason:
                            self._add_comment(paragraph, reason)
                    else:
                        break # 理論上ここには来ない

    def _insert_revision(self, paragraph, original, corrected):
        """変更履歴XML（w:del, w:ins）を生成して段落に追加します。"""
        timestamp = datetime.datetime.now().isoformat()
        
        # 削除タグ (w:del)
        w_del = OxmlElement('w:del')
        w_del.set(qn('w:id'), '1')
        w_del.set(qn('w:author'), 'AI Editor')
        w_del.set(qn('w:date'), timestamp)
        
        r_del = OxmlElement('w:r')
        delText = OxmlElement('w:delText')
        delText.text = original
        r_del.append(delText)
        w_del.append(r_del)
        paragraph._p.append(w_del)

        # 挿入タグ (w:ins)
        w_ins = OxmlElement('w:ins')
        w_ins.set(qn('w:id'), '2')
        w_ins.set(qn('w:author'), 'AI Editor')
        w_ins.set(qn('w:date'), timestamp)
        
        r_ins = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = corrected
        r_ins.append(t)
        w_ins.append(r_ins)
        paragraph._p.append(w_ins)

    def _add_comment(self, paragraph, reason):
        """段落にコメントを挿入します。"""
        try:
            # 最新のruns（挿入した部分など）に対してコメントを付ける
            self.doc.add_comment(paragraph.runs, text=reason, author="AI Editor", initials="AI")
        except Exception as e:
            print(f"コメント挿入エラー: {e}")

    def save(self, output_path: str):
        self.doc.save(output_path)
